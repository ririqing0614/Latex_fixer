import re
import os
import sys
from docx import Document
from docx.oxml import OxmlElement
from lxml import etree
import latex2mathml.converter
from docx.oxml import OxmlElement, parse_xml

# 获取当前脚本所在目录，确保离线 XSL 文件能被找到
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
XSL_PATH = os.path.join(BASE_DIR, "MML2OMML.XSL")

def latex_to_omml(latex_str):
    try:
        # 1. LaTeX -> MathML
        # 增加一些预处理，防止空内容或特殊转义字符导致崩溃
        if not latex_str.strip():
            return None
            
        mathml = latex2mathml.converter.convert(latex_str)
        
        if not os.path.exists(XSL_PATH):
            return "XSL_MISSING"
            
        xslt = etree.parse(XSL_PATH)
        transform = etree.XSLT(xslt)
        
        mathml_tree = etree.fromstring(mathml)
        omml_tree = transform(mathml_tree)
        
        return omml_tree.getroot()
        
    except Exception as e:
        #打印具体的公式内容和错误类型，方便排查
        print(f"  [DEBUG] 公式内容: {latex_str}")
        print(f"  [DEBUG] 错误详情: {type(e).__name__} - {e}")
        return None

def replace_latex_in_paragraph(paragraph):
    text = paragraph.text
    # 匹配 $...$ 或 $$...$$
    pattern = r'(\$\$.*?\$\$|\$.*?\$)'
    parts = re.split(pattern, text)
    
    if len(parts) <= 1:
        return

    paragraph.clear()
    
    for part in parts:
        if not part: continue
            
        if part.startswith('$') and part.endswith('$'):
            # 1. 预处理清理
            clean_latex = part.strip('$').strip()
            if not clean_latex:
                paragraph.add_run(part)
                continue

            # --- 2. 增强型预处理：解决 DoubleSuperscriptsError 和 积分方块问题 ---
            
            # (1) 处理连续撇号，防止解析器将其视为连续上标
            if "'''" in clean_latex:
                clean_latex = clean_latex.replace("'''", "^{\prime\prime\prime}")
            elif "''" in clean_latex:
                clean_latex = clean_latex.replace("''", "^{\prime\prime}")
            elif "'" in clean_latex:
                # 仅针对函数导数形式的单撇号处理
                clean_latex = re.sub(r"(?<=[a-zA-Z\)])'", r"^{\prime}", clean_latex)

            # (2) 修复微积分方块：确保积分号后有显式容器
            # 匹配 \int_{a}^{b} 或 \int_a^b 后的内容，并尝试用 {} 包裹
            if '\\int' in clean_latex:
                # 寻找积分限之后的被积函数位置，并用 {} 包裹剩余部分
                int_match = re.search(r'(\\int[_^0-9a-zA-Z{}]*)', clean_latex)
                if int_match:
                    prefix = int_match.group(0)
                    body = clean_latex[int_match.end():].strip()
                    if body and not body.startswith('{'):
                        clean_latex = f"{prefix} {{{body}}}"

            # --- 3. 执行转换 ---
            omml_element = latex_to_omml(clean_latex)
            
            if omml_element == "XSL_MISSING":
                paragraph.add_run(part)
            elif omml_element is not None:
                xml_str = etree.tostring(omml_element, encoding='unicode')
                try:
                    paragraph._element.append(parse_xml(xml_str))
                except Exception:
                    paragraph.add_run(part)
            else:
                # 如果转换依然失败，尝试最原始的 LaTeX 插入
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

def process_document(input_path, output_path):
    if not os.path.exists(XSL_PATH):
        print(f"!!! 错误: 在脚本目录下未找到 MML2OMML.XSL 文件 !!!")
        print(f"请从 Office 目录复制或从 W3C 下载该文件到: {XSL_PATH}")
        return

    print(f"正在加载文档: {input_path}")
    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"文件读取失败: {e}")
        return

    # 遍历所有段落
    for p in doc.paragraphs:
        if '$' in p.text:
            replace_latex_in_paragraph(p)

    # 遍历所有表格（针对实验题目中常见的函数表）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if '$' in p.text:
                        replace_latex_in_paragraph(p)

    try:
        doc.save(output_path)
        print(f"\n--- 处理成功 ---")
        print(f"输出文件: {output_path}")
    except Exception as e:
        print(f"文件保存失败: {e}")

if __name__ == "__main__":
    print("========================================")
    print("   Word LaTeX 自动化修复工具 (离线版)   ")
    print("========================================\n")
    
    in_path = input("1. 请粘贴输入 Word 文档的完整路径: ").strip('"').strip()
    if not in_path:
        print("路径不能为空。")
        sys.exit()

    out_path = input("2. 请输入输出路径 (直接回车保存在同目录): ").strip('"').strip()
    if not out_path:
        base, ext = os.path.splitext(in_path)
        out_path = f"{base}_公式版{ext}"

    process_document(in_path, out_path)