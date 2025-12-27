import streamlit as st
import os
import re
import io
from docx import Document
from docx.oxml import parse_xml
from lxml import etree
import latex2mathml.converter

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
XSL_PATH = os.path.join(BASE_DIR, "MML2OMML.XSL")

def latex_to_omml(latex_str):
    try:
        if not latex_str.strip(): return None
        # 预处理：修复连续撇号和积分号
        if "'''" in latex_str: latex_str = latex_str.replace("'''", "^{\prime\prime\prime}")
        elif "''" in latex_str: latex_str = latex_str.replace("''", "^{\prime\prime}")
        
        if '\\int' in latex_str and '{' not in latex_str:
            match = re.search(r'(\\int[_^0-9a-zA-Z]*)', latex_str)
            if match:
                prefix = match.group(0)
                body = latex_str[match.end():].strip()
                if body: latex_str = f"{prefix} {{{body}}}"

        mathml = latex2mathml.converter.convert(latex_str)
        if not os.path.exists(XSL_PATH): return "XSL_MISSING"
        
        xslt = etree.parse(XSL_PATH)
        transform = etree.XSLT(xslt)
        return transform(etree.fromstring(mathml)).getroot()
    except Exception:
        return None

def replace_latex_in_paragraph(paragraph):
    text = paragraph.text
    pattern = r'(\$\$.*?\$\$|\$.*?\$)'
    parts = re.split(pattern, text)
    if len(parts) <= 1: return

    paragraph.clear()
    for part in parts:
        if not part: continue
        if part.startswith('$') and part.endswith('$'):
            clean = part.strip('$').strip()
            omml = latex_to_omml(clean)
            if omml is not None and omml != "XSL_MISSING":
                xml_str = etree.tostring(omml, encoding='unicode')
                try:
                    paragraph._element.append(parse_xml(xml_str))
                except:
                    paragraph.add_run(part)
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

# --- 网页界面 ---
st.set_page_config(page_title="LaTeX 修复神器", page_icon="🧪")

st.sidebar.title("🛠️ 功能选择")
mode = st.sidebar.radio("请选择工作模式：", ["文档文件修复", "文本片段实时修复"])

if mode == "文档文件修复":
    st.title("📄 Word 文档公式修复")
    st.info("上传包含 LaTeX 乱码的 .docx 文件，系统将自动导出修复后的文档。")
    
    uploaded_file = st.file_uploader("选择 Word 文件", type=["docx"])
    
    if uploaded_file:
        doc = Document(uploaded_file)
        if st.button("开始修复文档"):
            with st.spinner('正在处理中...'):
                for p in doc.paragraphs:
                    if '$' in p.text: replace_latex_in_paragraph(p)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if '$' in p.text: replace_latex_in_paragraph(p)
            
            st.success("处理完成！")
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button("⬇️ 下载修复后的 Word", bio.getvalue(), f"Fixed_{uploaded_file.name}", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

else:
    st.title("📝 文本片段实时修复")
    st.markdown("在左侧输入包含 LaTeX 的文字，右侧将生成包含 **Word 可识别公式格式** 的预览。")
    st.caption("提示：由于网页无法直接显示 Word 公式对象，此处将尝试使用标准 LaTeX 预览，你可以检查公式结构是否正确。")

    col1, col2 = st.columns(2)
    
    with col1:
        input_text = st.text_area("输入区域 (粘贴文字片段):", height=300, placeholder="例如：已知 $f(x)=x^2$，求 $\int_0^1 f(x)dx$")
    
    with col2:
        st.write("预览输出 (渲染效果):")
        if input_text:
            st.markdown(input_text)
            
            with st.expander("查看转换后的 Word XML 代码 (供调试)"):
                test_match = re.search(r'\$(.*?)\$', input_text)
                if test_match:
                    res = latex_to_omml(test_match.group(1))
                    if res is not None and res != "XSL_MISSING":
                        st.code(etree.tostring(res, encoding='unicode', pretty_print=True), language='xml')
        else:
            st.gray("等待输入...")

st.sidebar.markdown("---")
st.sidebar.caption("Powered by Streamlit | ririqing0614")